#!/usr/bin/env python3
"""
extract_docx.py - Extract DOCX to Markdown chunks with format sidecars.

Produces chunk*.md and docx_structure.json used by build_docx.py.
All paragraphs share a single [P:N] index namespace across:
  - Document body (paragraphs + tables)
  - Section headers and footers
  - Footnotes and endnotes
"""

import argparse
import json
import os
import sys

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph as DocxParagraph
    from docx.table import Table as DocxTable
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

from manifest import create_manifest


_FOOTNOTES_REL = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
_ENDNOTES_REL = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes'
_SEPARATOR_IDS = {'-1', '0'}

_HF_TYPES = [
    ("header_default", lambda s: s.header),
    ("footer_default", lambda s: s.footer),
    ("header_first", lambda s: s.first_page_header),
    ("footer_first", lambda s: s.first_page_footer),
    ("header_even", lambda s: s.even_page_header),
    ("footer_even", lambda s: s.even_page_footer),
]

_HEADING_STYLES = {
    "heading 1": "#", "heading 2": "##", "heading 3": "###",
    "heading 4": "####", "heading 5": "#####",
}


# ---------------------------------------------------------------------------
# DOCX iteration
# ---------------------------------------------------------------------------

def iter_block_items(doc):
    """Yield Paragraph or Table objects from the document body in order."""
    body = doc.element.body
    for child in body:
        local = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if local == 'p':
            yield DocxParagraph(child, doc)
        elif local == 'tbl':
            yield DocxTable(child, doc)


# ---------------------------------------------------------------------------
# Format metadata helpers
# ---------------------------------------------------------------------------

def _pt(val):
    if val is None:
        return None
    try:
        return round(float(val.pt), 2)
    except (AttributeError, TypeError):
        return None


def _color_hex(color):
    try:
        rgb = color.rgb
        return str(rgb) if rgb else None
    except Exception:
        return None


def _run_meta(run):
    font = run.font
    return {
        "text": run.text,
        "bold": run.bold,
        "italic": run.italic,
        "underline": run.underline,
        "font_name": font.name,
        "font_size_pt": _pt(font.size),
        "color_rgb": _color_hex(font.color),
    }


def _alignment_str(para):
    try:
        return str(para.alignment).split('.')[-1] if para.alignment else None
    except Exception:
        return None


def _para_element(para, elem_idx):
    pf = para.paragraph_format
    return {
        "elem_idx": elem_idx,
        "type": "paragraph",
        "style": para.style.name if para.style else "Normal",
        "alignment": _alignment_str(para),
        "space_before_pt": _pt(pf.space_before),
        "space_after_pt": _pt(pf.space_after),
        "left_indent_pt": _pt(pf.left_indent),
        "runs": [_run_meta(r) for r in para.runs],
    }


# ---------------------------------------------------------------------------
# Inline markdown rendering
# ---------------------------------------------------------------------------

def _runs_to_inline_md(runs):
    """Build inline markdown text from a sequence of docx Run objects."""
    parts = []
    for run in runs:
        t = run.text
        if not t:
            continue
        if run.bold and run.italic:
            parts.append(f"***{t}***")
        elif run.bold:
            parts.append(f"**{t}**")
        elif run.italic:
            parts.append(f"*{t}*")
        else:
            parts.append(t)
    return "".join(parts)


def _runs_meta_to_inline_md(runs_meta):
    """Build inline markdown text from a list of run metadata dicts."""
    parts = []
    for run in runs_meta:
        t = run.get("text", "")
        if not t:
            continue
        bold = run.get("bold")
        italic = run.get("italic")
        if bold and italic:
            parts.append(f"***{t}***")
        elif bold:
            parts.append(f"**{t}**")
        elif italic:
            parts.append(f"*{t}*")
        else:
            parts.append(t)
    return "".join(parts)


def _para_to_md(para, elem_idx):
    text = _runs_to_inline_md(para.runs).strip()
    if not text:
        return ""
    style_key = (para.style.name or "").lower()
    prefix = _HEADING_STYLES.get(style_key, "")
    body = f"{prefix} {text}" if prefix else text
    return f"[P:{elem_idx}] {body}"


# ---------------------------------------------------------------------------
# Table element
# ---------------------------------------------------------------------------

def _table_element(table, base_idx):
    """Build table metadata. Deduplicates merged cells by XML element identity."""
    idx = base_idx
    cells_meta = []
    seen_tc = set()

    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            tc_id = id(cell._tc)
            is_dup = tc_id in seen_tc
            if not is_dup:
                seen_tc.add(tc_id)

            para_metas = []
            for para in cell.paragraphs:
                if not is_dup:
                    para_metas.append({
                        "elem_idx": idx,
                        "type": "cell_paragraph",
                        "runs": [_run_meta(r) for r in para.runs],
                    })
                idx += 1  # Always increment to keep shared namespace in sync

            cells_meta.append({
                "row": ri, "col": ci,
                "paragraphs": para_metas,
                "is_merged_duplicate": is_dup,
            })

    return {
        "elem_idx": base_idx,
        "type": "table",
        "rows": len(table.rows),
        "cols": len(table.columns),
        "cells": cells_meta,
    }, idx


def _table_to_md(table, cells_meta):
    """Render table as Markdown with inline formatting and [P:N] markers."""
    if not table.rows:
        return ""

    cell_map = {}
    for cell_data in cells_meta:
        r, c = cell_data["row"], cell_data["col"]
        if cell_data.get("is_merged_duplicate"):
            cell_map[(r, c)] = (None, "")
            continue
        paras = cell_data["paragraphs"]
        if paras:
            idx = paras[0]["elem_idx"]
            text = " ".join(
                _runs_meta_to_inline_md(p["runs"]).strip() for p in paras
            ).strip()
            cell_map[(r, c)] = (idx, text)
        else:
            cell_map[(r, c)] = (None, "")

    lines = []
    ncols = len(table.columns)
    for ri in range(len(table.rows)):
        cells = []
        for ci in range(ncols):
            idx, text = cell_map.get((ri, ci), (None, ""))
            cells.append(f"[P:{idx}] {text}" if idx is not None else text)
        lines.append("| " + " | ".join(cells) + " |")
        if ri == 0:
            lines.append("| " + " | ".join(["---"] * ncols) + " |")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Headers and footers
# ---------------------------------------------------------------------------

def _extract_section_headers_footers(doc, base_idx):
    """Extract header/footer paragraphs from all sections.

    Returns (items, all_elements, next_idx).
    """
    items = []
    all_elements = []
    idx = base_idx

    for si, section in enumerate(doc.sections):
        for hf_type_name, getter in _HF_TYPES:
            try:
                hf = getter(section)
                if hf is None or hf.is_linked_to_previous:
                    continue
            except Exception:
                continue

            for pi, para in enumerate(hf.paragraphs):
                text = _runs_to_inline_md(para.runs).strip()
                meta = {
                    "elem_idx": idx,
                    "type": "hf_paragraph",
                    "section_idx": si,
                    "hf_type": hf_type_name,
                    "para_idx_within_hf": pi,
                    "runs": [_run_meta(r) for r in para.runs],
                }
                all_elements.append(meta)
                if text:
                    items.append((f"[P:{idx}] {text}", meta))
                idx += 1

    return items, all_elements, idx


# ---------------------------------------------------------------------------
# Footnotes and endnotes
# ---------------------------------------------------------------------------

def _get_notes_elem(doc, rel_type):
    """Return the notes part XML element, or None."""
    try:
        for rel in doc.part.rels.values():
            if rel.reltype == rel_type:
                return rel.target_part._element
    except Exception:
        pass
    return None


def _extract_notes(doc, base_idx, rel_type, note_tag, note_type_name):
    """Extract footnote or endnote paragraphs. Returns (items, all_elements, next_idx)."""
    items = []
    all_elements = []
    idx = base_idx

    notes_elem = _get_notes_elem(doc, rel_type)
    if notes_elem is None:
        return items, all_elements, idx

    for note_elem in notes_elem.findall(qn(note_tag)):
        note_id = note_elem.get(qn('w:id'), '')
        if note_id in _SEPARATOR_IDS:
            continue
        for pi, p_elem in enumerate(note_elem.findall(qn('w:p'))):
            para = DocxParagraph(p_elem, doc)
            text = _runs_to_inline_md(para.runs).strip()
            meta = {
                "elem_idx": idx,
                "type": note_type_name,
                "note_id": note_id,
                "para_idx_within_note": pi,
                "runs": [_run_meta(r) for r in para.runs],
            }
            all_elements.append(meta)
            if text:
                items.append((f"[P:{idx}] {text}", meta))
            idx += 1

    return items, all_elements, idx


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------

def _group_into_chunks(items, target_size):
    chunks, current, size = [], [], 0
    for md, meta in items:
        block_size = len(md) + 2
        if size + block_size > target_size and current:
            chunks.append(current)
            current, size = [], 0
        current.append((md, meta))
        size += block_size
    if current:
        chunks.append(current)
    return chunks


# ---------------------------------------------------------------------------
# Main extraction
# ---------------------------------------------------------------------------

def extract_docx_to_chunks(docx_path, temp_dir, chunk_size=6000):
    """Extract DOCX to Markdown chunks + docx_structure.json.

    Returns number of chunk files created, or 0 on failure.
    """
    doc = Document(docx_path)

    items = []
    all_elements = []
    elem_idx = 0

    # Body
    for block in iter_block_items(doc):
        if isinstance(block, DocxParagraph):
            md = _para_to_md(block, elem_idx)
            meta = _para_element(block, elem_idx)
            items.append((md, meta))
            all_elements.append(meta)
            elem_idx += 1
        elif isinstance(block, DocxTable):
            table_meta, elem_idx = _table_element(block, elem_idx)
            md = _table_to_md(block, table_meta["cells"])
            items.append((md, table_meta))
            all_elements.append(table_meta)

    # Headers and footers
    hf_items, hf_elements, elem_idx = _extract_section_headers_footers(doc, elem_idx)
    items.extend(hf_items)
    all_elements.extend(hf_elements)

    # Footnotes
    fn_items, fn_elements, elem_idx = _extract_notes(
        doc, elem_idx, _FOOTNOTES_REL, 'w:footnote', 'footnote_paragraph'
    )
    items.extend(fn_items)
    all_elements.extend(fn_elements)

    # Endnotes
    en_items, en_elements, elem_idx = _extract_notes(
        doc, elem_idx, _ENDNOTES_REL, 'w:endnote', 'endnote_paragraph'
    )
    items.extend(en_items)
    all_elements.extend(en_elements)

    structure = {
        "schema_version": 2,
        "source_docx": os.path.abspath(docx_path),
        "elements": all_elements,
    }
    struct_path = os.path.join(temp_dir, "docx_structure.json")
    with open(struct_path, "w", encoding="utf-8") as f:
        json.dump(structure, f, ensure_ascii=False, indent=2)
    print(f"Wrote docx_structure.json ({elem_idx} elements)")

    translatable = [(md, meta) for md, meta in items if md.strip()]
    groups = _group_into_chunks(translatable, chunk_size)
    if not groups:
        print("No translatable content found in DOCX")
        return 0

    input_md = os.path.join(temp_dir, "input.md")
    with open(input_md, "w", encoding="utf-8") as f:
        f.write("\n\n".join(md for md, _ in translatable))

    chunk_files = []
    for i, group in enumerate(groups, 1):
        fname = f"chunk{i:04d}.md"
        with open(os.path.join(temp_dir, fname), "w", encoding="utf-8") as f:
            f.write("\n\n".join(md for md, _ in group))
        chunk_files.append(fname)

    create_manifest(temp_dir, chunk_files, input_md)
    print(f"DOCX extracted: {len(chunk_files)} chunks, {elem_idx} total elements")
    return len(chunk_files)


def main():
    parser = argparse.ArgumentParser(description="Extract DOCX to Markdown chunks")
    parser.add_argument("docx_file", help="Input DOCX file")
    parser.add_argument("--temp-dir", required=True, help="Output temp directory")
    parser.add_argument("--chunk-size", type=int, default=6000)
    args = parser.parse_args()

    if not os.path.exists(args.docx_file):
        print(f"Error: {args.docx_file} not found")
        sys.exit(1)

    os.makedirs(args.temp_dir, exist_ok=True)
    count = extract_docx_to_chunks(args.docx_file, args.temp_dir, args.chunk_size)
    if count == 0:
        sys.exit(1)
    print(f"Done. Chunks in: {args.temp_dir}")


if __name__ == "__main__":
    main()
