#!/usr/bin/env python3
"""
build_docx.py - Rebuild DOCX from translated Markdown chunks + docx_structure.json.

Maps [P:N] markers in translated output.md back to original DOCX formatting.
Handles body, section headers/footers, footnotes, and endnotes.
Inline markdown (***bold+italic***, **bold**, *italic*) is parsed back to runs.
"""

import argparse
import json
import os
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph as DocxParagraph
    from docx.table import Table as DocxTable
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


_MARKER_RE = re.compile(r'^\[P:(\d+)\]\s*')
_INLINE_MD_RE = re.compile(r'\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*')

_FOOTNOTES_REL = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
_ENDNOTES_REL = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes'
_SEPARATOR_IDS = {'-1', '0'}

_HF_TYPES = [
    ("header_default", lambda s: s.header),
    ("footer_default", lambda s: s.footer),
    ("header_first", lambda s: s.first_page_header),
    ("footer_first", lambda s: s.first_page_footer),
    ("header_even", lambda s: s.even_page_header),
    ("footer_even", lambda s: s.even_page_footer),
]


# ---------------------------------------------------------------------------
# Load structure index
# ---------------------------------------------------------------------------

def load_structure(temp_dir):
    """Load docx_structure.json. Returns (elem_index, source_docx_path)."""
    path = os.path.join(temp_dir, "docx_structure.json")
    if not os.path.exists(path):
        return None, None
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)

    index = {}
    for elem in data.get("elements", []):
        etype = elem.get("type")
        if etype in ("paragraph", "cell_paragraph", "hf_paragraph",
                     "footnote_paragraph", "endnote_paragraph"):
            index[elem["elem_idx"]] = elem
        elif etype == "table":
            for cell in elem.get("cells", []):
                for para in cell.get("paragraphs", []):
                    index[para["elem_idx"]] = para

    return index, data.get("source_docx")


# ---------------------------------------------------------------------------
# Parse translations from output.md
# ---------------------------------------------------------------------------

def parse_translations(output_md):
    """Return {elem_idx: translated_text} from [P:N] markers in output.md."""
    translations = {}
    with open(output_md, "r", encoding="utf-8") as f:
        content = f.read()

    for line in content.splitlines():
        stripped = line.strip().strip("|").strip()
        m = _MARKER_RE.match(stripped)
        if not m:
            continue
        idx = int(m.group(1))
        text = stripped[m.end():].strip()
        text = re.sub(r'^#{1,6}\s+', '', text)
        translations[idx] = text

    return translations


# ---------------------------------------------------------------------------
# Inline markdown → run segments
# ---------------------------------------------------------------------------

def _parse_inline_md(text):
    """Parse ***x***, **x**, *x* into [(text, bold, italic)] segments.

    bold/italic are True/False when explicitly set by a marker, None to inherit.
    """
    segments = []
    last_end = 0
    for m in _INLINE_MD_RE.finditer(text):
        if m.start() > last_end:
            segments.append((text[last_end:m.start()], None, None))
        if m.group(1) is not None:    # ***bold+italic***
            segments.append((m.group(1), True, True))
        elif m.group(2) is not None:  # **bold**
            segments.append((m.group(2), True, None))
        else:                          # *italic*
            segments.append((m.group(3), None, True))
        last_end = m.end()
    if last_end < len(text):
        segments.append((text[last_end:], None, None))
    return segments or [(text, None, None)]


# ---------------------------------------------------------------------------
# Apply translated text to a paragraph
# ---------------------------------------------------------------------------

def _first_text_run(para):
    for r in para.runs:
        if r.text:
            return r
    return None


def _apply_run_format(run, meta):
    if meta.get("bold") is not None:
        run.bold = meta["bold"]
    if meta.get("italic") is not None:
        run.italic = meta["italic"]
    if meta.get("underline") is not None:
        run.underline = meta["underline"]
    if meta.get("font_name"):
        run.font.name = meta["font_name"]
    if meta.get("font_size_pt"):
        run.font.size = Pt(meta["font_size_pt"])
    if meta.get("color_rgb"):
        try:
            h = meta["color_rgb"].lstrip("#")
            run.font.color.rgb = RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
        except Exception:
            pass


def set_para_text(para, text, elem_meta):
    """Replace paragraph text, restoring inline bold/italic from markdown markers."""
    if not text:
        return

    original_runs = elem_meta.get("runs", [])
    base_meta = original_runs[0] if original_runs else {}

    for r in para.runs:
        r.text = ""

    segments = _parse_inline_md(text)
    has_inline = any(bold is not None or italic is not None for _, bold, italic in segments)

    if not has_inline:
        first_run = _first_text_run(para)
        if first_run is not None:
            first_run.text = text
            if base_meta:
                _apply_run_format(first_run, base_meta)
        else:
            new_run = para.add_run(text)
            if base_meta:
                _apply_run_format(new_run, base_meta)
        return

    first_run = _first_text_run(para)
    for i, (seg_text, seg_bold, seg_italic) in enumerate(segments):
        run = first_run if (i == 0 and first_run is not None) else para.add_run("")
        run.text = seg_text
        if base_meta:
            _apply_run_format(run, base_meta)
        if seg_bold is not None:
            run.bold = seg_bold
        if seg_italic is not None:
            run.italic = seg_italic


# ---------------------------------------------------------------------------
# Rebuild body
# ---------------------------------------------------------------------------

def _iter_body(doc):
    for child in doc.element.body:
        local = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if local in ('p', 'tbl'):
            yield local, child


def _rebuild_body(doc, translations, elem_index):
    elem_idx = 0
    for tag, xml_elem in _iter_body(doc):
        if tag == 'p':
            para = DocxParagraph(xml_elem, doc)
            if elem_idx in translations and elem_idx in elem_index:
                set_para_text(para, translations[elem_idx], elem_index[elem_idx])
            elem_idx += 1
        elif tag == 'tbl':
            table = DocxTable(xml_elem, doc)
            seen_tc = set()
            for row in table.rows:
                for cell in row.cells:
                    tc_id = id(cell._tc)
                    is_dup = tc_id in seen_tc
                    if not is_dup:
                        seen_tc.add(tc_id)
                    for para in cell.paragraphs:
                        if not is_dup and elem_idx in translations and elem_idx in elem_index:
                            set_para_text(para, translations[elem_idx], elem_index[elem_idx])
                        elem_idx += 1


# ---------------------------------------------------------------------------
# Rebuild headers and footers
# ---------------------------------------------------------------------------

def _rebuild_headers_footers(doc, elem_index, translations):
    hf_lookup = {
        (m["section_idx"], m["hf_type"], m["para_idx_within_hf"]): eidx
        for eidx, m in elem_index.items()
        if m.get("type") == "hf_paragraph"
    }
    if not hf_lookup:
        return

    for si, section in enumerate(doc.sections):
        for hf_type_name, getter in _HF_TYPES:
            try:
                hf = getter(section)
                if hf is None or hf.is_linked_to_previous:
                    continue
            except Exception:
                continue
            for pi, para in enumerate(hf.paragraphs):
                eidx = hf_lookup.get((si, hf_type_name, pi))
                if eidx is not None and eidx in translations:
                    set_para_text(para, translations[eidx], elem_index[eidx])


# ---------------------------------------------------------------------------
# Rebuild footnotes and endnotes
# ---------------------------------------------------------------------------

def _get_notes_part_elem(doc, rel_type):
    try:
        for rel in doc.part.rels.values():
            if rel.reltype == rel_type:
                return rel.target_part._element
    except Exception:
        pass
    return None


def _rebuild_notes(doc, elem_index, translations, rel_type, note_tag, note_type_name):
    notes_lookup = {
        (m["note_id"], m["para_idx_within_note"]): eidx
        for eidx, m in elem_index.items()
        if m.get("type") == note_type_name
    }
    if not notes_lookup:
        return

    notes_elem = _get_notes_part_elem(doc, rel_type)
    if notes_elem is None:
        return

    for note_elem in notes_elem.findall(qn(note_tag)):
        note_id = note_elem.get(qn('w:id'), '')
        if note_id in _SEPARATOR_IDS:
            continue
        for pi, p_elem in enumerate(note_elem.findall(qn('w:p'))):
            eidx = notes_lookup.get((note_id, pi))
            if eidx is not None and eidx in translations:
                para = DocxParagraph(p_elem, doc)
                set_para_text(para, translations[eidx], elem_index[eidx])


# ---------------------------------------------------------------------------
# Main rebuild
# ---------------------------------------------------------------------------

def rebuild_docx(source_docx, temp_dir, output_docx):
    """Rebuild DOCX by injecting translated text into a copy of source_docx."""
    elem_index, _ = load_structure(temp_dir)
    if elem_index is None:
        print("Error: docx_structure.json not found in temp dir")
        return False

    output_md = os.path.join(temp_dir, "output.md")
    if not os.path.exists(output_md):
        print("Error: output.md not found")
        return False

    translations = parse_translations(output_md)
    if not translations:
        print("Warning: no [P:N] markers found in output.md — nothing to rebuild")
        return False

    doc = Document(source_docx)
    _rebuild_body(doc, translations, elem_index)
    _rebuild_headers_footers(doc, elem_index, translations)
    _rebuild_notes(doc, elem_index, translations, _FOOTNOTES_REL, 'w:footnote', 'footnote_paragraph')
    _rebuild_notes(doc, elem_index, translations, _ENDNOTES_REL, 'w:endnote', 'endnote_paragraph')

    doc.save(output_docx)
    size = os.path.getsize(output_docx)
    print(f"Native DOCX rebuilt: {output_docx} ({size:,} bytes)")
    return True


def main():
    parser = argparse.ArgumentParser(description="Rebuild DOCX from translated chunks")
    parser.add_argument("source_docx", help="Original DOCX (used as formatting template)")
    parser.add_argument("--temp-dir", required=True, help="Temp directory with output.md + docx_structure.json")
    parser.add_argument("--output", required=True, help="Output DOCX path")
    args = parser.parse_args()

    if not os.path.exists(args.source_docx):
        print(f"Error: {args.source_docx} not found")
        sys.exit(1)

    if not rebuild_docx(args.source_docx, args.temp_dir, args.output):
        sys.exit(1)


if __name__ == "__main__":
    main()
